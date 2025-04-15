from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches
import tempfile
import os

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        data = request.form
        logo_file = request.files["logo"]
        logo_path = None

        if logo_file and logo_file.filename:
            logo_path = os.path.join(app.config['UPLOAD_FOLDER'], logo_file.filename)
            logo_file.save(logo_path)

        doc = Document("template.docx")
        replacements = {
            "[Client Company Name]": data["company_name"],
            "[Date]": data["effective_date"],
            "[Version Number]": data["version"],
            "John Doe": data["author"],
            "Jane Smith": data["reviewer"],
            "CEO": data["approver"],
            "security@acmetech.com": data["email"],
            "555-123-4567": data["phone"],
            "Microsoft Teams": data["im_platform"],
            "email and secure employee portal": data["distribution"]
        }

        for para in doc.paragraphs:
            for key, val in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, val)

        if logo_path:
            doc.sections[0].header.paragraphs[0].text = ""
            doc.sections[0].header.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.5))

        doc.add_page_break()
        doc.add_heading("Signature Acknowledgment", level=1)
        doc.add_paragraph("Employee Name: ______________________________")
        doc.add_paragraph("Signature: _________________________________")
        doc.add_paragraph("Date: ______________________________________")

        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_docx.name)
        return send_file(temp_docx.name, as_attachment=True, download_name="Custom_AUP.docx")

    return render_template("form.html")

if __name__ == "__main__":
    app.run(debug=True)
